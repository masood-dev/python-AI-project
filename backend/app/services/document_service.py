"""Document parsing and storage service."""
import io
from bson import ObjectId
import pdfplumber
from docx import Document as DocxDocument
from app.models.document import Document
from app.utils.exceptions import DocumentParsingError, InvalidFileTypeError, FileTooLargeError
from app.core.config import settings
from app.db.mongodb import get_database


async def parse_and_store_document(
    session_id: ObjectId,
    filename: str,
    file_bytes: bytes,
    file_type: str,
) -> Document:
    """
    Parse uploaded document and store in MongoDB.

    Args:
        session_id: Session ID
        filename: Original filename
        file_bytes: File content bytes
        file_type: File extension (pdf, docx, txt)

    Returns:
        Stored Document object

    Raises:
        InvalidFileTypeError: If file type not supported
        FileTooLargeError: If file exceeds size limit
        DocumentParsingError: If parsing fails
    """
    # Validate file type
    if file_type.lower() not in settings.allowed_file_types:
        raise InvalidFileTypeError(file_type, settings.allowed_file_types)

    # Validate file size
    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > settings.max_file_size_mb:
        raise FileTooLargeError(int(size_mb), settings.max_file_size_mb)

    # Parse based on file type
    try:
        if file_type.lower() == "pdf":
            extracted_text = _parse_pdf(file_bytes)
        elif file_type.lower() == "docx":
            extracted_text = _parse_docx(file_bytes)
        elif file_type.lower() == "txt":
            extracted_text = _parse_txt(file_bytes)
        else:
            raise DocumentParsingError(filename, "Unknown file type")
    except DocumentParsingError:
        raise
    except Exception as e:
        raise DocumentParsingError(filename, str(e))

    # Store in MongoDB
    document = Document(
        session_id=session_id,
        filename=filename,
        file_type=file_type.lower(),
        extracted_text=extracted_text,
    )

    db = get_database()
    result = await db.documents.insert_one(document.to_dict())
    document._id = result.inserted_id

    return document


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF."""
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
    except Exception as e:
        raise DocumentParsingError("PDF", str(e))

    if not text_parts:
        raise DocumentParsingError("PDF", "No extractable text found (may be scanned image)")

    return "\n".join(text_parts)


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX."""
    text_parts = []
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
    except Exception as e:
        raise DocumentParsingError("DOCX", str(e))

    if not text_parts:
        raise DocumentParsingError("DOCX", "No extractable text found")

    return "\n".join(text_parts)


def _parse_txt(file_bytes: bytes) -> str:
    """Extract text from TXT."""
    try:
        text = file_bytes.decode("utf-8")
        if not text.strip():
            raise DocumentParsingError("TXT", "File is empty")
        return text
    except UnicodeDecodeError as e:
        raise DocumentParsingError("TXT", f"Invalid UTF-8 encoding: {str(e)}")
    except Exception as e:
        raise DocumentParsingError("TXT", str(e))
